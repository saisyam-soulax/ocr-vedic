from __future__ import annotations

import asyncio
import json
import logging
import os
import shutil
import time
from contextlib import asynccontextmanager
from datetime import UTC, datetime
from typing import TYPE_CHECKING, Annotated

import starlette.requests
from fastapi import Depends, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response, StreamingResponse

if TYPE_CHECKING:
    from starlette.datastructures import FormData

# ---------------------------------------------------------------------------
# Root logger
# ---------------------------------------------------------------------------
_log_level = os.environ.get("LOG_LEVEL", "INFO").upper()
logging.basicConfig(
    level=getattr(logging, _log_level, logging.INFO),
    format="%(asctime)s %(levelname)s %(name)s: %(message)s",
)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Raise Starlette's per-part size limit so large PDFs aren't silently rejected.
# ---------------------------------------------------------------------------
_MAX_PART_SIZE = 500 * 1024 * 1024  # 500 MiB
_original_get_form = starlette.requests.Request._get_form


async def _patched_get_form(
    self: starlette.requests.Request,
    *,
    max_files: int | float = 1000,
    max_fields: int | float = 1000,
    max_part_size: int = _MAX_PART_SIZE,
) -> "FormData":
    return await _original_get_form(
        self,
        max_files=max_files,
        max_fields=max_fields,
        max_part_size=max_part_size,
    )


starlette.requests.Request._get_form = _patched_get_form  # type: ignore[method-assign]

from app import vllm_runtime
from app.config import Settings, get_settings
import re as _re

from app.schemas import (
    ErrorResponse,
    HealthResponse,
    OcrDefaultsResponse,
    OcrJobResponse,
    OcrProvider,
    OcrSavedResult,
    ProvidersResponse,
    ProviderInfo,
    RecentJobSummary,
    VllmState,
    VllmStatusResponse,
)
from app.services.ocr_service import (
    build_few_shots_for_provider,
    get_default_system_prompt,
    parse_few_shots_json,
    run_ocr_job,
)
from app.storage_uploads import persist_ocr_uploads, prune_old_batches, write_batch_metadata

# ---------------------------------------------------------------------------
# In-memory job registry  { job_id -> (queue, background_task, created_at) }
# ---------------------------------------------------------------------------
_jobs: dict[str, tuple[asyncio.Queue, asyncio.Task | None, float]] = {}
_JOB_TTL = 1200  # 20 minutes — prune abandoned jobs


def _prune_stale_jobs() -> None:
    now = time.time()
    stale = [k for k, (_, task, t) in _jobs.items() if now - t > _JOB_TTL]
    for k in stale:
        _, task, _ = _jobs.pop(k)
        if task and not task.done():
            task.cancel()
        logger.info("Pruned stale job: %s", k)


# ---------------------------------------------------------------------------
# Provider configuration check (no blocking HTTP calls here)
# ---------------------------------------------------------------------------

def _provider_configured(settings: Settings, p: OcrProvider) -> tuple[bool, str | None]:
    if p == OcrProvider.gemini:
        if settings.gemini_use_vertexai:
            ok = bool(settings.google_api_key) or bool(settings.google_cloud_project)
            return ok, None if ok else (
                "Vertex AI mode requires GOOGLE_API_KEY or GOOGLE_CLOUD_PROJECT + ADC"
            )
        ok = bool(settings.google_api_key)
        return ok, None if ok else "Set GOOGLE_API_KEY"
    if p == OcrProvider.bedrock_claude:
        ok = bool(settings.aws_region and settings.bedrock_claude_model_id)
        return ok, None if ok else "Set AWS_REGION and BEDROCK_CLAUDE_MODEL_ID"
    if p == OcrProvider.bedrock_ocr:
        ok = bool(settings.aws_region and settings.bedrock_ocr_model_id)
        return ok, None if ok else "Set AWS_REGION and BEDROCK_OCR_MODEL_ID"
    # vllm_gemma — just check config, not live reachability (use /api/vllm/status for that)
    if not settings.vllm_enabled:
        return False, "Set VLLM_ENABLED=true"
    if not settings.vllm_base_url:
        return False, "Set VLLM_BASE_URL"
    return True, None


# ---------------------------------------------------------------------------
# App factory
# ---------------------------------------------------------------------------

@asynccontextmanager
async def lifespan(app: FastAPI):
    settings = get_settings()
    logger.info(
        "App startup: log_level=%s concurrency=%d dpi=%d",
        _log_level, settings.ocr_page_concurrency, settings.ocr_pdf_dpi,
    )
    yield
    logger.info("App shutdown")


def create_app() -> FastAPI:
    app = FastAPI(
        title="Vedic OCR API",
        version="0.2.0",
        lifespan=lifespan,
        description="Multimodal OCR for Devanāgarī + IAST with streaming page output.",
    )
    settings = get_settings()
    origins = [o.strip() for o in settings.cors_origins.split(",") if o.strip()]
    app.add_middleware(
        CORSMiddleware,
        allow_origins=origins or ["*"],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

    # ------------------------------------------------------------------
    # Health
    # ------------------------------------------------------------------

    @app.get("/health", response_model=HealthResponse)
    def health() -> HealthResponse:
        return HealthResponse(status="ok")

    # ------------------------------------------------------------------
    # Providers
    # ------------------------------------------------------------------

    @app.get("/api/providers", response_model=ProvidersResponse)
    def providers(s: Settings = Depends(get_settings)) -> ProvidersResponse:
        labels = {
            OcrProvider.gemini: "Google Gemini",
            OcrProvider.bedrock_claude: "AWS Bedrock — Claude",
            OcrProvider.bedrock_ocr: "AWS Bedrock — Open multimodal",
            OcrProvider.vllm_gemma: "Local — Gemma 4 (vLLM)",
        }
        rows: list[ProviderInfo] = []
        for p in OcrProvider:
            configured, detail = _provider_configured(s, p)
            if p == OcrProvider.gemini:
                default_mid, mids = s.gemini_models_for_providers()
            elif p == OcrProvider.bedrock_claude:
                default_mid, mids = s.bedrock_claude_models_for_providers()
            elif p == OcrProvider.bedrock_ocr:
                default_mid, mids = s.bedrock_open_models_for_providers()
            else:
                default_mid, mids = s.vllm_models_for_providers()
            rows.append(
                ProviderInfo(
                    id=p.value,
                    label=labels[p],
                    configured=configured,
                    detail=detail if not configured else None,
                    default_model_id=default_mid,
                    model_options=mids,
                )
            )
        return ProvidersResponse(providers=rows)

    @app.get("/api/ocr/defaults", response_model=OcrDefaultsResponse)
    def ocr_defaults() -> OcrDefaultsResponse:
        return OcrDefaultsResponse(system_prompt=get_default_system_prompt())

    # ------------------------------------------------------------------
    # vLLM lifecycle
    # ------------------------------------------------------------------

    @app.get("/api/vllm/status", response_model=VllmStatusResponse)
    async def vllm_status(s: Settings = Depends(get_settings)) -> VllmStatusResponse:
        if not s.vllm_enabled:
            return VllmStatusResponse(
                state=VllmState.stopped, reachable=False, message="VLLM_ENABLED=false"
            )

        internal = vllm_runtime.get_state()

        # Active lifecycle states always win — avoids the polling resetting the UI
        # back to "stopped" while a load/unload is in progress.
        if internal in (VllmState.starting, VllmState.ready, VllmState.stopping):
            return VllmStatusResponse(
                state=internal,
                reachable=internal == VllmState.ready,
                message=vllm_runtime.get_error(),
            )
        if internal == VllmState.error:
            return VllmStatusResponse(
                state=VllmState.error, reachable=False, message=vllm_runtime.get_error()
            )

        # State is stopped — do a live health-check so we pick up an
        # externally-started vLLM (e.g. the user ran docker compose manually).
        reachable = await vllm_runtime.check_health(s.vllm_base_url)
        if reachable:
            # External vLLM is up — sync our internal state.
            await vllm_runtime.load(s)
            return VllmStatusResponse(state=VllmState.ready, reachable=True)
        return VllmStatusResponse(state=VllmState.stopped, reachable=False)

    @app.post("/api/vllm/load", response_model=VllmStatusResponse)
    async def vllm_load(s: Settings = Depends(get_settings)) -> VllmStatusResponse:
        if not s.vllm_enabled:
            raise HTTPException(400, "VLLM_ENABLED=false — enable vLLM in .env first")
        await vllm_runtime.load(s)
        state = vllm_runtime.get_state()
        return VllmStatusResponse(
            state=state,
            reachable=state == VllmState.ready,
            message=vllm_runtime.get_error(),
        )

    @app.post("/api/vllm/unload", response_model=VllmStatusResponse)
    async def vllm_unload(s: Settings = Depends(get_settings)) -> VllmStatusResponse:
        if not s.vllm_enabled:
            raise HTTPException(400, "VLLM_ENABLED=false")
        await vllm_runtime.unload(s)
        return VllmStatusResponse(state=VllmState.stopped, reachable=False)

    # ------------------------------------------------------------------
    # OCR — submit job (returns immediately)
    # ------------------------------------------------------------------

    @app.post("/api/ocr", response_model=OcrJobResponse)
    async def submit_ocr(
        files: Annotated[list[UploadFile], File(description="PDF and/or image files")],
        provider: str = Form(...),
        system_prompt: str | None = Form(None),
        few_shots: str | None = Form(None),
        few_shot_files: Annotated[
            list[UploadFile] | None,
            File(description="Images for few-shots; same order as few_shots"),
        ] = None,
        model_id: str | None = Form(None),
        s: Settings = Depends(get_settings),
    ) -> OcrJobResponse:
        _prune_stale_jobs()

        logger.info(
            "OCR submit: provider=%s files=%d model=%s",
            provider, len(files), model_id or "(default)",
        )

        # Validate provider
        try:
            prov = OcrProvider(provider)
        except ValueError:
            raise HTTPException(
                422,
                detail=f"Invalid provider '{provider}'. Use: gemini, bedrock_claude, bedrock_ocr, vllm_gemma.",
            )
        ok, msg = _provider_configured(s, prov)
        if not ok:
            raise HTTPException(400, detail=msg or "Provider not configured")

        model_id_clean: str | None = None
        if model_id is not None:
            stripped = model_id.strip()
            if not stripped:
                raise HTTPException(422, detail="model_id must be non-empty when supplied")
            model_id_clean = stripped

        # Build few-shots
        parsed_shots = parse_few_shots_json(few_shots)
        try:
            resolved_shots = await build_few_shots_for_provider(parsed_shots, few_shot_files or None)
        except HTTPException:
            raise

        # Persist uploads to disk
        upload_root = s.upload_root_path()
        upload_root.mkdir(parents=True, exist_ok=True)
        if s.upload_retain_hours > 0:
            prune_old_batches(upload_root, s.upload_retain_hours)

        batch_id, batch_dir, saved = await persist_ocr_uploads(upload_root, files)
        write_batch_metadata(
            batch_dir,
            {
                "batch_id": batch_id,
                "created_at": datetime.now(UTC).isoformat(),
                "provider": provider,
                "model_id": model_id_clean,
                "files": [fn for (_, fn, _) in saved],
                "few_shot_count": len(resolved_shots),
            },
        )

        # Create job queue and start background processing.
        q: asyncio.Queue = asyncio.Queue()
        task = asyncio.create_task(
            run_ocr_job(
                queue=q,
                saved_files=saved,
                provider=provider,
                system_prompt=system_prompt,
                few_shots=resolved_shots,
                settings=s,
                model_id=model_id_clean,
                batch_dir=batch_dir,
                retain=s.upload_retain,
            )
        )
        job_id = batch_id
        _jobs[job_id] = (q, task, time.time())

        return OcrJobResponse(
            job_id=job_id,
            stream_url=f"/api/ocr/{job_id}/stream",
            total_files=len(saved),
        )

    # ------------------------------------------------------------------
    # OCR — SSE stream (long-lived GET)
    # ------------------------------------------------------------------

    @app.get("/api/ocr/{job_id}/stream")
    async def stream_ocr_results(job_id: str) -> StreamingResponse:
        entry = _jobs.get(job_id)
        if not entry:
            raise HTTPException(404, detail="Job not found. It may have expired or already been consumed.")
        q, task, _ = entry

        async def generate():
            try:
                while True:
                    try:
                        item = await asyncio.wait_for(q.get(), timeout=25.0)
                    except asyncio.TimeoutError:
                        # Keep-alive comment so nginx/browser don't close the connection.
                        yield ": ping\n\n"
                        continue
                    if item is None:
                        # Job complete — close the stream.
                        return
                    event_name = item.get("event", "message")
                    yield f"event: {event_name}\ndata: {json.dumps(item)}\n\n"
            finally:
                _jobs.pop(job_id, None)
                if task and not task.done():
                    task.cancel()

        return StreamingResponse(
            generate(),
            media_type="text/event-stream",
            headers={
                # Tells nginx (and any other buffering proxy) not to buffer this response.
                "X-Accel-Buffering": "no",
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
            },
        )

    # ------------------------------------------------------------------
    # OCR — retrieve saved results (survives browser refresh / disconnect)
    # ------------------------------------------------------------------

    _UUID_RE = _re.compile(r"^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$")

    def _read_saved_result(batch_dir, job_id: str) -> OcrSavedResult:
        results_path = batch_dir / "results.jsonl"
        pages = []
        if results_path.exists():
            for raw in results_path.read_text(encoding="utf-8").splitlines():
                if raw.strip():
                    try:
                        pages.append(json.loads(raw))
                    except json.JSONDecodeError:
                        pass
        pages.sort(key=lambda p: p.get("index", 0))

        complete_path = batch_dir / "job_complete.json"
        done = complete_path.exists()
        total = len(pages)
        elapsed = None
        completed_at = None
        if done:
            try:
                info = json.loads(complete_path.read_text(encoding="utf-8"))
                total = info.get("total", len(pages))
                elapsed = info.get("elapsed_seconds")
                completed_at = info.get("completed_at")
            except (json.JSONDecodeError, OSError):
                pass

        provider = None
        files: list[str] = []
        submitted_at = None
        meta_path = batch_dir / "metadata.json"
        if meta_path.exists():
            try:
                meta = json.loads(meta_path.read_text(encoding="utf-8"))
                provider = meta.get("provider")
                files = meta.get("files", [])
                submitted_at = meta.get("created_at")
            except (json.JSONDecodeError, OSError):
                pass

        return OcrSavedResult(
            job_id=job_id,
            done=done,
            total=total,
            done_count=len(pages),
            provider=provider,
            files=files,
            submitted_at=submitted_at,
            completed_at=completed_at,
            elapsed_seconds=elapsed,
            pages=pages,
        )

    @app.get("/api/ocr/jobs", response_model=dict)
    async def list_ocr_jobs(s: Settings = Depends(get_settings)) -> dict:
        """List recent OCR jobs (completed or in-progress) from disk."""
        upload_root = s.upload_root_path()
        jobs: list[RecentJobSummary] = []
        if upload_root.is_dir():
            entries = sorted(
                (e for e in upload_root.iterdir() if e.is_dir()),
                key=lambda e: e.stat().st_mtime,
                reverse=True,
            )
            for batch_dir in entries[:30]:
                meta_path = batch_dir / "metadata.json"
                if not meta_path.exists():
                    continue
                try:
                    saved = _read_saved_result(batch_dir, batch_dir.name)
                    jobs.append(RecentJobSummary(
                        job_id=saved.job_id,
                        provider=saved.provider,
                        files=saved.files,
                        submitted_at=saved.submitted_at,
                        done=saved.done,
                        total=saved.total,
                        done_count=saved.done_count,
                        completed_at=saved.completed_at,
                    ))
                except Exception:
                    pass
        return {"jobs": [j.model_dump() for j in jobs]}

    @app.get("/api/ocr/{job_id}/result", response_model=OcrSavedResult)
    async def get_ocr_result(job_id: str, s: Settings = Depends(get_settings)) -> OcrSavedResult:
        """Return saved pages for a completed (or in-progress) job."""
        if not _UUID_RE.match(job_id):
            raise HTTPException(400, "Invalid job ID format")
        batch_dir = s.upload_root_path() / job_id
        if not batch_dir.is_dir():
            raise HTTPException(404, "Job not found or results have expired")
        return _read_saved_result(batch_dir, job_id)

    def _build_combined_text(saved: OcrSavedResult) -> str:
        parts = []
        for p in saved.pages:
            header = f"## {p.source_file}"
            if p.page_in_source is not None:
                header += f" (page {p.page_in_source})"
            parts.append(f"{header}\n\n{p.text}")
        return "\n\n---\n\n".join(parts)

    @app.get("/api/ocr/{job_id}/download.txt")
    async def download_txt(job_id: str, s: Settings = Depends(get_settings)) -> Response:
        """Download all saved pages as a plain-text file."""
        if not _UUID_RE.match(job_id):
            raise HTTPException(400, "Invalid job ID format")
        batch_dir = s.upload_root_path() / job_id
        if not batch_dir.is_dir():
            raise HTTPException(404, "Job not found or results have expired")
        saved = _read_saved_result(batch_dir, job_id)
        if not saved.pages:
            raise HTTPException(404, "No pages have been saved for this job yet")
        content = _build_combined_text(saved)
        slug = (saved.files[0].rsplit(".", 1)[0] if saved.files else job_id[:8])
        return Response(
            content=content.encode("utf-8"),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="vedic-ocr-{slug}.txt"'},
        )

    @app.get("/api/ocr/{job_id}/download.docx")
    async def download_docx(job_id: str, s: Settings = Depends(get_settings)) -> Response:
        """Download all saved pages as a Word document."""
        import io
        from docx import Document
        from docx.shared import Pt

        if not _UUID_RE.match(job_id):
            raise HTTPException(400, "Invalid job ID format")
        batch_dir = s.upload_root_path() / job_id
        if not batch_dir.is_dir():
            raise HTTPException(404, "Job not found or results have expired")
        saved = _read_saved_result(batch_dir, job_id)
        if not saved.pages:
            raise HTTPException(404, "No pages have been saved for this job yet")

        doc = Document()
        # Set default font to Noto Serif for Devanagari rendering
        style = doc.styles["Normal"]
        style.font.name = "Noto Serif"
        style.font.size = Pt(11)

        for p in saved.pages:
            heading = f"{p.source_file}"
            if p.page_in_source is not None:
                heading += f" (page {p.page_in_source})"
            h = doc.add_heading(heading, level=2)
            h.style.font.name = "Noto Serif"
            for line in p.text.split("\n"):
                para = doc.add_paragraph(line)
                para.style.font.name = "Noto Serif"
            doc.add_paragraph("―" * 40)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        slug = (saved.files[0].rsplit(".", 1)[0] if saved.files else job_id[:8])
        return Response(
            content=buf.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="vedic-ocr-{slug}.docx"'},
        )

    # ------------------------------------------------------------------
    # Error handlers
    # ------------------------------------------------------------------

    @app.exception_handler(HTTPException)
    async def http_exc_handler(request, exc: HTTPException):
        return JSONResponse(
            status_code=exc.status_code,
            content=ErrorResponse(detail=str(exc.detail), code="http_error").model_dump(),
        )

    return app


app = create_app()
